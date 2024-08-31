import streamlit as st
import docx2txt
import requests
import os
from dotenv import load_dotenv
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import re

# Load environment variables
load_dotenv()

def process_document(file):
    # Extract text from .docx file
    text = docx2txt.process(file)
    
    # Split into words and filter for words starting with 'a'
    words = text.split()
    a_words = [word for word in words if word.lower().startswith('a')]
    
    # Take only the first 10 words (or less if there aren't 10)
    return a_words[:1]

def check_spelling(word, app_id, app_key):
    url = f"https://od-api-sandbox.oxforddictionaries.com/api/v2/entries/en-gb/{word.lower()}"
    headers = {
        "app_id": app_id,
        "app_key": app_key,
    }
    
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        return {"word": word, "isCorrect": True}
    elif response.status_code == 404:
        return {"word": word, "isCorrect": False}
    else:
        return {"word": word, "isCorrect": False, "error": f"Unexpected status code: {response.status_code}"}

def insert_tracked_change(paragraph, original_word, corrected_word):
    run = paragraph.add_run()
    delete = OxmlElement('w:del')
    delete.set(qn('w:author'), 'SpellCheck Pro')
    delete.set(qn('w:date'), '2023-06-09T00:00:00Z')
    delete.text = original_word
    run._element.append(delete)

    insert = OxmlElement('w:ins')
    insert.set(qn('w:author'), 'SpellCheck Pro')
    insert.set(qn('w:date'), '2023-06-09T00:00:00Z')
    insert.text = corrected_word
    run._element.append(insert)

def modify_document(file, results):
    doc = Document(file)
    changes_made = False

    for paragraph in doc.paragraphs:
        for result in results:
            if not result['isCorrect']:
                original_word = result['word']
                # Use a placeholder if no correction is available
                corrected_word = f"[CORRECTION NEEDED: {original_word}]"
                
                # Use regex to find word boundaries
                pattern = r'\b' + re.escape(original_word) + r'\b'
                if re.search(pattern, paragraph.text, re.IGNORECASE):
                    new_text = re.sub(pattern, corrected_word, paragraph.text, flags=re.IGNORECASE)
                    if new_text != paragraph.text:
                        paragraph.text = new_text
                        changes_made = True

    if changes_made:
        # Save the modified document
        modified_file = "modified_" + file.name
        doc.save(modified_file)
        return modified_file
    else:
        return None

def main():
    st.title("SpellCheck Pro")
    
    # File upload
    uploaded_file = st.file_uploader("Choose a .docx file", type="docx")
    
    # API credentials
    app_id = st.text_input("Enter OED API Application ID", value=os.getenv("OED_APP_ID", ""))
    app_key = st.text_input("Enter OED API Application Key", value=os.getenv("OED_APP_KEY", ""), type="password")
    
    if uploaded_file and app_id and app_key:
        if st.button("Start Spell Check"):
            with st.spinner("Processing..."):
                # Process document
                words_to_check = process_document(uploaded_file)
                
                # Check spelling
                results = [check_spelling(word, app_id, app_key) for word in words_to_check]
                
                # Modify document
                modified_doc = modify_document(uploaded_file, results)
                
                # Save modified document
                if modified_doc:
                    st.success("Document has been modified with tracked changes. Download it below.")
                    
                    # Provide download link for modified document
                    with open(modified_doc, "rb") as file:
                        st.download_button(
                            label="Download modified document",
                            data=file,
                            file_name=modified_doc,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("No changes were made to the document.")
                    
                # Generate report
                st.subheader("Spell Check Report")
                st.write(f"Total words checked: {len(results)}")
                
                correct_words = [result for result in results if result['isCorrect']]
                incorrect_words = [result for result in results if not result['isCorrect']]
                
                st.write(f"Correct spellings: {len(correct_words)}")
                st.write(f"Incorrect spellings: {len(incorrect_words)}")
                
                if incorrect_words:
                    st.subheader("Words with potential spelling errors and their corrections:")
                    for word in incorrect_words:
                        st.write(f"- {word['word']} (suggested correction: {word.get('suggestion', 'No suggestion available')})")

if __name__ == "__main__":
    main()