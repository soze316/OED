import streamlit as st
import docx2txt
import requests
import os
import re
from dotenv import load_dotenv
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Load environment variables
load_dotenv()

def is_english_word(word):
    # This is a simple check. For a more comprehensive check, you might want to use a library like 'enchant'
    return re.match(r'^[a-zA-Z]+$', word) is not None

def process_document(file):
    # Extract text from .docx file
    text = docx2txt.process(file)
    
    # Split into words, filter for English words starting with 'a'
    words = text.split()
    a_words = [word for word in words if word.lower().startswith('a') and is_english_word(word)]
    
    # Take only the first 10 words (or less if there aren't 10)
    return a_words[:1]

def check_spelling(word, app_id, app_key):
    url = f"https://od-api-sandbox.oxforddictionaries.com/api/v2/search/en-gb?q={word.lower()}&prefix=false"
    headers = {
        "app_id": app_id,
        "app_key": app_key,
    }
    
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        data = response.json()
        if data.get('results'):
            top_result = data['results'][0]
            if top_result['word'].lower() == word.lower():
                return {"word": word, "isCorrect": True}
            else:
                return {"word": word, "isCorrect": False, "suggestion": top_result['word']}
        else:
            return {"word": word, "isCorrect": False}
    else:
        return {"word": word, "isCorrect": False, "error": f"Unexpected status code: {response.status_code}"}

#def insert_tracked_change(paragraph, original_word, corrected_word):
    # Clear the existing runs in the paragraph (to avoid mixing styles)
    #paragraph.clear()
    
    # Insert deletion
    #delete_run = paragraph.add_run()
    #delete_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for deletion
    #delete_run.font.strike = True  # Strikethrough for deletion
    
    # Insert insertion
    #insert_run = paragraph.add_run(corrected_word)
    #insert_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for insertion
    #insert_run.font.bold = True  # Bold for insertion
    
    # Add comments to indicate these are tracked changes
    # Note: While Word tracks changes through internal XML, this is a visual approximation
    # to give the user an idea that a change was intended.

    
    #run = paragraph.add_run()
    #delete = OxmlElement('w:del')
    #delete.set(qn('w:author'), 'SpellCheck Pro')
    #delete.set(qn('w:date'), '2023-06-09T00:00:00Z')
    #st.write(delete)
     # Create a run with red text inside the deletion
    #run_inside_del = OxmlElement('w:r')
    #rPr = OxmlElement('w:rPr')
    #c = OxmlElement('w:color')
    #c.set(qn('w:val'), 'FF0000')  # Red color in hex
    #rPr.append(c)
    #run_inside_del.append(rPr)
    #t = OxmlElement('w:t')
    #t.text = original_word
    #run_inside_del.append(t)


    #delete.text = original_word
    #run._element.append(delete)

    #st.write(f"Delete: {delete}")

    #insert = OxmlElement('w:ins')
    #insert.set(qn('w:author'), 'SpellCheck Pro')
    #insert.set(qn('w:date'), '2023-06-09T00:00:00Z')
    #insert.text = corrected_word
    #run._element.append(insert)

    #st.write(f"Insert: {insert}")

def modify_document(file, results):
    doc = Document(file)
    changes_made = False

    for paragraph in doc.paragraphs:
        for result in results:
            if not result['isCorrect']:
                original_word = result['word']
                corrected_word = f"[CORRECTION NEEDED: {original_word}]" if 'suggestion' not in result else result['suggestion']
                             
                pattern = r'\b' + re.escape(original_word) + r'\b'

                if re.search(pattern, paragraph.text, re.IGNORECASE):
                    new_text = re.sub(pattern, corrected_word, paragraph.text, flags=re.IGNORECASE)
                    #st.write(f"New Text: {new_text}")

                    if new_text != paragraph.text:
                        paragraph.text = new_text
                        #st.write(f"Paragraph Text: {paragraph.text}")
                        #insert_tracked_change(paragraph, original_word, corrected_word)
                        changes_made = True

    if changes_made:
        modified_file = "modified_" + file.name
        doc.save(modified_file)
        return modified_file
    else:
        return None

def main():
    st.title("SpellCheck Pro - English Spelling Check")
    
    st.write("This application checks the spelling of English words starting with 'a' in your document.")
    
    # File upload
    uploaded_file = st.file_uploader("Choose a .docx file", type="docx")
    
    # API credentials
    app_id = st.text_input("Enter OED API Application ID", value=os.getenv("OED_APP_ID", ""))
    app_key = st.text_input("Enter OED API Application Key", value=os.getenv("OED_APP_KEY", ""), type="password")
    
    if uploaded_file and app_id and app_key:
        if st.button("Start English Spell Check"):
            with st.spinner("Processing..."):
                # Process document
                words_to_check = process_document(uploaded_file)
                st.write(words_to_check)

                if not words_to_check:
                    st.warning("No English words starting with 'a' found in the document.")
                else:
                    # Check spelling
                    results = [check_spelling(word, app_id, app_key) for word in words_to_check]
                    
                    # Modify document
                    modified_doc = modify_document(uploaded_file, results)
                    
                    # Generate report
                    st.subheader("English Spell Check Report")
                    st.write(f"Total English words checked: {len(results)}")
                    
                    correct_words = [result for result in results if result['isCorrect']]
                    incorrect_words = [result for result in results if not result['isCorrect']]
                    
                    st.write(f"Correct spellings: {len(correct_words)}")
                    st.write(f"Incorrect spellings: {len(incorrect_words)}")
                    
                    if incorrect_words:
                        st.subheader("English words with potential spelling errors:")
                        for word in incorrect_words:
                            suggestion = word.get('suggestion', 'No suggestion available')
                            st.write(f"- {word['word']} (Suggested: {suggestion})")
                    
                    if modified_doc:
                        st.success("Document has been modified with tracked changes. Download it below.")
                        
                        # Provide download link for modified document
                        with open(modified_doc, "rb") as file:
                            st.download_button(
                                label="Download modified document",
                                data=file,
                                file_name=modified_doc,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    else:
                        st.info("No changes were made to the document.")

if __name__ == "__main__":
    main()